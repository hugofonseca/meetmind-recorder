import discord
from discord.ext import commands, voice_recv
import datetime
import os
import wave
import asyncio
import threading
import io
from faster_whisper import WhisperModel
from docx import Document
from collections import defaultdict
from dotenv import load_dotenv
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import logging
import openai
from openai import AsyncOpenAI
import json
import pickle
import time # Added for background processor
from typing import Optional
import queue
import numpy as np
import zipfile




logging.basicConfig(level=logging.INFO)
logging.getLogger("discord.ext.voice_recv").setLevel(logging.INFO)
logging.getLogger("discord.voice_state").setLevel(logging.INFO)
logger = logging.getLogger(__name__)


load_dotenv()


intents = discord.Intents.default()
intents.message_content = True
intents.voice_states = True
intents.members = True

bot = commands.Bot(command_prefix="!", intents=intents)



openai_client = None
openai_key = os.getenv("OPENAI_API_KEY")

if openai_key:
    openai_client = AsyncOpenAI(api_key=openai_key)
    print("✅ OpenAI habilitado (Q&A disponível).")
else:
    print("ℹ️ OpenAI desabilitado (Q&A indisponível). Rode sem OPENAI_API_KEY.")



model = WhisperModel("small", device="cpu", compute_type="int8")


SUPPORTED_LANGUAGES = {
   
    "english": "en", "en": "en",
    "spanish": "es", "es": "es", "español": "es",
    "french": "fr", "fr": "fr", "français": "fr",
    "german": "de", "de": "de", "deutsch": "de",
    "italian": "it", "it": "it", "italiano": "it",
    "portuguese": "pt", "pt": "pt", "português": "pt",
    "russian": "ru", "ru": "ru", "русский": "ru",
    "chinese": "zh", "zh": "zh", "中文": "zh", "mandarin": "zh",
    "japanese": "ja", "ja": "ja", "日本語": "ja",
    "korean": "ko", "ko": "ko", "한국어": "ko",
    "hindi": "hi", "hi": "hi", "हिंदी": "hi",
    
    
    "arabic": "ar", "ar": "ar", "العربية": "ar",
    
    
    "dutch": "nl", "nl": "nl", "nederlands": "nl",
    "swedish": "sv", "sv": "sv", "svenska": "sv",
    "norwegian": "no", "no": "no", "norsk": "no",
    "danish": "da", "da": "da", "dansk": "da",
    "finnish": "fi", "fi": "fi", "suomi": "fi",
    "polish": "pl", "pl": "pl", "polski": "pl",
    "czech": "cs", "cs": "cs", "čeština": "cs",
    "hungarian": "hu", "hu": "hu", "magyar": "hu",
    "greek": "el", "el": "el", "ελληνικά": "el",
    "turkish": "tr", "tr": "tr", "türkçe": "tr",
    "hebrew": "he", "he": "he", "עברית": "he",
    
   
    "thai": "th", "th": "th", "ไทย": "th",
    "vietnamese": "vi", "vi": "vi", "tiếng việt": "vi",
    "indonesian": "id", "id": "id", "bahasa indonesia": "id",
    "malay": "ms", "ms": "ms", "bahasa melayu": "ms",
    "tamil": "ta", "ta": "ta", "தமிழ்": "ta",
    "bengali": "bn", "bn": "bn", "বাংলা": "bn",
    "urdu": "ur", "ur": "ur", "اردو": "ur",
    "persian": "fa", "fa": "fa", "فارسی": "fa", "farsi": "fa",
    
   
    "auto": "auto", "automatic": "auto", "detect": "auto"
}

def get_language_code(language_input):
    """Convert language input to standardized language code."""
    if not language_input:
        return "en"
    
    lang_lower = language_input.lower().strip()
    return SUPPORTED_LANGUAGES.get(lang_lower, "en")

def get_language_display_name(language_code):
    """Get display name for language code."""
    lang_names = {
        "en": "English", "es": "Spanish", "fr": "French", "de": "German",
        "it": "Italian", "pt": "Portuguese", "ru": "Russian", "zh": "Chinese",
        "ja": "Japanese", "ko": "Korean", "hi": "Hindi", "ar": "Arabic",
        "nl": "Dutch", "sv": "Swedish", "no": "Norwegian", "da": "Danish",
        "fi": "Finnish", "pl": "Polish", "cs": "Czech", "hu": "Hungarian",
        "el": "Greek", "tr": "Turkish", "he": "Hebrew", "th": "Thai",
        "vi": "Vietnamese", "id": "Indonesian", "ms": "Malay", "ta": "Tamil",
        "bn": "Bengali", "ur": "Urdu", "fa": "Persian", "auto": "Auto-detect"
    }
    return lang_names.get(language_code, language_code.upper())


active_meetings = {}

def save_meetings():
    """Save active meetings to persistent storage."""
    try:
        # Convert meetings to serializable format
        serializable_meetings = {}
        for guild_id, meeting in active_meetings.items():
            serializable_meeting = {
                "channel_id": meeting["channel"].id if meeting["channel"] else None,
                "guild_id": guild_id,
                "log": meeting["log"],
                "start_time": meeting["start_time"],
                "started_by": meeting["started_by"],
                "language": meeting["language"],
                "language_display": meeting["language_display"]
            }
            serializable_meetings[guild_id] = serializable_meeting
        
        with open("meetings.pkl", "wb") as f:
            pickle.dump(serializable_meetings, f)
        logger.info(f"Saved {len(serializable_meetings)} meetings to persistent storage")
    except Exception as e:
        logger.error(f"Error saving meetings: {e}")

def load_meetings():
    """Load meetings from persistent storage."""
    try:
        if os.path.exists("meetings.pkl"):
            with open("meetings.pkl", "rb") as f:
                saved_meetings = pickle.load(f)
            
            # Convert back to active format (without voice client references)
            for guild_id, saved_meeting in saved_meetings.items():
                # Check if meeting is still valid (within last 24 hours)
                if (datetime.datetime.now() - saved_meeting["start_time"]).total_seconds() < 86400:
                    active_meetings[guild_id] = {
                        "channel": None,  # Will be restored when bot connects
                        "log": saved_meeting["log"],
                        "vc": None,  # Will be restored when bot connects
                        "start_time": saved_meeting["start_time"],
                        "started_by": saved_meeting["started_by"],
                        "language": saved_meeting["language"],
                        "language_display": saved_meeting["language_display"],
                        "sink": None
                    }
                    logger.info(f"Restored meeting for guild {guild_id}")
                else:
                    logger.info(f"Meeting for guild {guild_id} expired, not restoring")
            
            logger.info(f"Loaded {len(active_meetings)} meetings from persistent storage")
    except Exception as e:
        logger.error(f"Error loading meetings: {e}")


class MixRecorder:
    """
    Grava um WAV 'mix' (somando áudio de múltiplos usuários).
    - write() só coloca frames numa fila (rápido, thread-safe)
    - um worker thread alinha frames por 'tick' (20ms) e mistura com numpy
    """

    def __init__(self, meeting, out_dir="meeting_audio", sample_rate=48000, channels=2, sampwidth=2, frame_ms=20):
        self.meeting = meeting
        self.sample_rate = sample_rate
        self.channels = channels
        self.sampwidth = sampwidth
        self.frame_ms = frame_ms

        os.makedirs(out_dir, exist_ok=True)
        start_tag = meeting["start_time"].strftime("%Y%m%d_%H%M%S")
        self.filepath = os.path.join(out_dir, f"meeting_mix_{start_tag}_{meeting['channel'].guild.id}.wav")

        self._wf = wave.open(self.filepath, "wb")
        self._wf.setnchannels(channels)
        self._wf.setsampwidth(sampwidth)
        self._wf.setframerate(sample_rate)

        self._q = queue.Queue(maxsize=2000)
        self._stop = threading.Event()
        self._worker = threading.Thread(target=self._run, daemon=True)

        # Para alinhar frames aproximando para "ticks" de 20ms
        self._t0 = time.monotonic()
        self._tick_sec = frame_ms / 1000.0

        # Buffer: tick_index -> {uid: pcm_bytes}
        self._bucket = {}
        self._last_flushed = -1

        self._worker.start()

    def push_frame(self, uid: int, pcm_bytes: bytes):
        """Enfileira um frame PCM do usuário."""
        if self._stop.is_set():
            return
        try:
            now = time.monotonic()
            tick = int((now - self._t0) / self._tick_sec)
            self._q.put_nowait((tick, uid, pcm_bytes))
        except queue.Full:
            # Se lotar, dropa frame para não travar (melhor do que travar o bot)
            pass

    def close(self):
        """Encerra thread e fecha o arquivo WAV."""
        self._stop.set()
        try:
            self._worker.join(timeout=2.0)
        except Exception:
            pass
        try:
            self._wf.close()
        except Exception:
            pass

    def _mix_frames(self, frames: list[bytes]) -> bytes:
        """
        Mix simples:
        - converte para int16
        - soma com clamp (evita clipping)
        """
        if not frames:
            # silêncio: 20ms stereo 48kHz => 960 samples/ch => 960*2ch = 1920 samples => *2 bytes = 3840 bytes
            samples_per_ch = int(self.sample_rate * (self.frame_ms / 1000.0))
            total_samples = samples_per_ch * self.channels
            return (np.zeros(total_samples, dtype=np.int16)).tobytes()

        arrays = [np.frombuffer(f, dtype=np.int16) for f in frames]
        mix = np.sum(arrays, axis=0)

        # clamp
        mix = np.clip(mix, -32768, 32767).astype(np.int16)
        return mix.tobytes()

    def _flush_tick(self, tick: int):
        """
        Flush do tick: mistura todos frames disponíveis desse tick e grava no WAV.
        """
        per_user = self._bucket.pop(tick, {})
        mixed = self._mix_frames(list(per_user.values()))
        self._wf.writeframes(mixed)

    def _run(self):
        """
        Worker: coleta frames, agrupa por tick e flush com pequeno atraso
        (para permitir receber frames de múltiplos usuários do mesmo tick).
        """
        # atraso para "esperar" frames daquele tick chegarem
        max_lag_ticks = 3  # ~60ms

        while not self._stop.is_set():
            try:
                tick, uid, pcm = self._q.get(timeout=0.1)
                self._bucket.setdefault(tick, {})[uid] = pcm

                # Flush ticks antigos
                flush_upto = tick - max_lag_ticks
                while self._last_flushed < flush_upto:
                    self._last_flushed += 1
                    self._flush_tick(self._last_flushed)

            except queue.Empty:
                # Em idle, flush alguns ticks se houver pendência antiga
                if self._bucket:
                    oldest = min(self._bucket.keys())
                    # se o oldest já é bem antigo, flush
                    now_tick = int((time.monotonic() - self._t0) / self._tick_sec)
                    if now_tick - oldest > max_lag_ticks:
                        while self._last_flushed < oldest:
                            self._last_flushed += 1
                            self._flush_tick(self._last_flushed)

        # Ao fechar, flush tudo que restou em ordem
        if self._bucket:
            for t in sorted(self._bucket.keys()):
                self._flush_tick(t)
            self._bucket.clear()






async def fail_meeting_capture(guild_id: int, meeting: dict, reason: str, ctx: Optional[commands.Context] = None):
    """
    Encerra a reunião como NÃO CAPTURADA (fallback).
    - Interrompe listener/sink
    - Desconecta do voice
    - Notifica no canal de transcrição e/ou canal do comando
    - Remove do active_meetings e persiste (sem manter como meeting válida ativa)
    """
    try:
        # Marca status (não persiste como "meeting válida", mas útil para logs locais)
        meeting["capture_failed"] = True
        meeting["capture_fail_reason"] = reason

        # 1) Interromper sink/listener
        sink = meeting.get("sink")
        if sink and hasattr(sink, "cleanup"):
            sink.cleanup()

        vc = meeting.get("vc")
        if vc:
            # alguns builds expõem stop_listening/is_listening; se existir, use
            if hasattr(vc, "is_listening") and vc.is_listening():
                if hasattr(vc, "stop_listening"):
                    try:
                        vc.stop_listening()
                    except Exception:
                        pass

            # 2) Desconectar da voz
            try:
                await vc.disconnect()
            except Exception:
                pass

        # 3) Mensagem para o usuário (canal de transcrição e canal do comando)
        stage_msg = (
            "⚠️ **Falha na captura de áudio neste canal de voz (stream corrompido/Opus).**\n"
            "✅ Para transcrição confiável, use o Stage: **📌 Meeting Room (Transcrição)**\n"
            "➡️ Entre no Stage e rode novamente: `!start_meeting english`\n"
        )

        # Canal principal de aviso: transcription channel, se existir
        ch = meeting.get("channel")
        if ch:
            try:
                await ch.send(stage_msg)
            except Exception:
                pass

        # Também avisa no canal onde o comando foi executado (se disponível)
        if ctx and ctx.channel and (not ch or ctx.channel.id != ch.id):
            try:
                await ctx.send(stage_msg)
            except Exception:
                pass

    finally:
        # 4) Remover do active_meetings e salvar estado (não fica reunião ativa “válida”)
        if guild_id in active_meetings:
            active_meetings.pop(guild_id, None)
        save_meetings()



class TranscriptionSink(voice_recv.AudioSink):
    def __init__(self, meeting, bot_instance, language="auto"):
        super().__init__()
        self.stopped = False  # When True, no further processing or sending occurs
        self._bg_stop_event = threading.Event()  # Signal to stop background processor
        self.meeting = meeting
        self.bot = bot_instance
        self.language = language
        self.buffers = defaultdict(bytearray)
        self.last_time = defaultdict(lambda: datetime.datetime.now())
        self.processing = defaultdict(bool)
        self.pending_transcriptions = []  # Queue for pending transcriptions
        self.processing_lock = asyncio.Lock()  # Lock for sequential processing
        self.next_sequence = 0  # Sequence number for ordering
        
         # === MixRecorder (áudio mix) ===
        os.makedirs("meeting_audio", exist_ok=True)
        self.mix_path = os.path.join(
            "meeting_audio",
            f"meeting_mix_{self.meeting['start_time'].strftime('%Y%m%d_%H%M%S')}.wav"
        )

        self.mix_wav = wave.open(self.mix_path, "wb")
        self.mix_wav.setnchannels(2)
        self.mix_wav.setsampwidth(2)
        self.mix_wav.setframerate(48000)

        self.meeting["mix_audio_path"] = self.mix_path

        logger.info(f"[MixRecorder] WAV mix criado em: {self.mix_path}")

        # Audio processing configuration - Optimized for low latency
        self.buffer_duration = 4.0  # Reduced from 5.0s for faster response
        self.min_buffer_size = 16000  # Reduced for faster processing
        self.max_buffer_size = 24000  # Reduced to prevent delays
        self.force_process_interval = 1.5  # Force process every 1.5s if no natural triggers
        
        # Real-time mode settings
        self.real_time_mode = True  # Enable real-time processing
        self.streaming_mode = True  # Enable streaming transcriptions

        self.filepath = self.mix_path  # compatibilidade: alguns trechos podem esperar .filepath
        
        
        
        # Start background processing timer
        self.start_background_processor()
    
    def start_background_processor(self):
        """Start background timer to force process audio at regular intervals."""
        def background_processor():
            while not self._bg_stop_event.is_set():
                try:
                    time.sleep(self.force_process_interval)
                    # Force process any pending audio
                    loop = self.bot.loop
                    if loop and not loop.is_closed():
                        asyncio.run_coroutine_threadsafe(
                            self.force_process_pending_audio(), 
                            loop
                        )
                except Exception as e:
                    logger.error(f"Background processor error: {e}")
                    break
        
        threading.Thread(target=background_processor, daemon=True).start()
    
    async def force_process_pending_audio(self):
        """Force process any audio that has been waiting too long."""
        if self.stopped:
            return
        current_time = datetime.datetime.now()
        
        for uid, buffer in self.buffers.items():
            if len(buffer) > 0 and not self.processing[uid]:
                elapsed = (current_time - self.last_time[uid]).total_seconds()
                
                # Force process if buffer has been waiting too long
                if elapsed >= self.force_process_interval:
                    await self.process_user_audio(uid, current_time)
    
    async def process_user_audio(self, uid, current_time):
        """Process audio for a specific user."""
        if self.stopped:
            return
        if self.processing[uid] or len(self.buffers[uid]) < self.min_buffer_size:
            return
        
        # Get user object
        user = None
        for guild in self.bot.guilds:
            member = guild.get_member(uid)
            if member:
                user = member
                break
        
        if not user:
            return
        
        # Process the audio
        pcm_data = bytes(self.buffers[uid])
        self.buffers[uid] = bytearray()
        self.last_time[uid] = current_time
        self.processing[uid] = True
        
        # Create transcription task
        transcription_task = {
            'user': user,
            'pcm_data': pcm_data,
            'timestamp': current_time,
            'sequence': self.next_sequence,
            'uid': uid
        }
        self.next_sequence += 1
        
        # Process immediately
        await self.process_audio_async(transcription_task)
        
    def wants_opus(self) -> bool:
        return False
    
    def cleanup(self):
        # Signal all processors to stop and clear state
        self.stopped = True

        # bg stop event pode não existir se init falhar
        try:
            if hasattr(self, "_bg_stop_event") and self._bg_stop_event:
                self._bg_stop_event.set()
        except Exception:
            pass
        
        # fecha wav se existir
        try:
            if hasattr(self, "mix_wav") and self.mix_wav:
                self.mix_wav.close()
                logger.info(f"[MixRecorder] WAV mix fechado: {getattr(self, 'mix_path', 'sem caminho')}")
        except Exception as e:
            logger.error(f"[MixRecorder] Erro ao fechar WAV mix: {e}")

        # limpa estruturas (se existirem)
        try:
            self.buffers.clear()
            self.last_time.clear()
            self.processing.clear()
            self.pending_transcriptions.clear()
        except Exception:
            pass


    def write(self, user, data: voice_recv.VoiceData):
        """Called by voice_recv when audio arrives (must be sync)."""
        if self.stopped:
            return
                
        if not data.pcm or not user:
            return
                    
        # grava o mix (simples: append do PCM recebido)
        try:
            self.mix_wav.writeframes(data.pcm)
        except Exception as e:
            logger.error(f"[MixRecorder] Erro ao gravar WAV mix: {e}")
        
        uid = user.id
        
        # Add audio to user's buffer
        self.buffers[uid].extend(data.pcm)
        
        now = datetime.datetime.now()
        elapsed = (now - self.last_time[uid]).total_seconds()
        current_buffer_size = len(self.buffers[uid])
        
        # Process audio when buffer is ready (time-based or size-based)
        should_process = (
            elapsed >= self.buffer_duration and 
            current_buffer_size >= self.min_buffer_size and 
            not self.processing[uid]
        )
        
        # Force processing if buffer gets too large (prevents long delays)
        if current_buffer_size >= self.max_buffer_size and not self.processing[uid]:
            should_process = True
        
        if should_process:
            # Schedule processing (não bloquear a thread do voice_recv)
            def schedule_transcription():
                try:
                    loop = self.bot.loop
                    if loop and not loop.is_closed():
                        asyncio.run_coroutine_threadsafe(
                            self.process_user_audio(uid, now),
                            loop
                        )
                except Exception as e:
                    logger.error(f"Error scheduling transcription: {e}")
                    self.processing[uid] = False

            threading.Thread(target=schedule_transcription, daemon=True).start()


    
    async def process_audio_async(self, transcription_task):
        """Process audio data asynchronously with proper ordering."""
        if self.stopped:
            return
        user = transcription_task['user']
        pcm_data = transcription_task['pcm_data']
        uid = transcription_task['uid']
        timestamp = transcription_task['timestamp']
        sequence = transcription_task['sequence']
        
        try:
            # Convert audio to WAV format
            wav_buffer = io.BytesIO()
            with wave.open(wav_buffer, 'wb') as wav_file:
                wav_file.setnchannels(2)  # Stereo
                wav_file.setsampwidth(2)  # 16-bit
                wav_file.setframerate(48000)  # 48kHz
                wav_file.writeframes(pcm_data)
            
            wav_buffer.seek(0)
            
            # Transcribe audio
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(None, self.transcribe_audio, wav_buffer)
            
            if self.stopped:
                return
            if text and text.strip():
                # Add to pending transcriptions queue
                async with self.processing_lock:
                    if self.stopped:
                        return
                    self.pending_transcriptions.append({
                        'timestamp': timestamp,
                        'sequence': sequence,
                        'speaker': user.display_name,
                        'text': text.strip(),
                        'user_id': uid
                    })
                    
                    # Sort by timestamp and sequence to maintain order
                    self.pending_transcriptions.sort(key=lambda x: (x['timestamp'], x['sequence']))
                    
                    # Process all pending transcriptions in order
                    await self.process_pending_transcriptions()
                    
        except Exception as e:
            logger.error(f"Error processing audio for user {user.display_name}: {e}")
        finally:
            self.processing[uid] = False
    
    async def process_pending_transcriptions(self):
        """Process all pending transcriptions in chronological order."""
        if self.stopped:
            # Drop anything queued without sending
            self.pending_transcriptions.clear()
            return
        while self.pending_transcriptions and not self.stopped:
            # Get the next transcription in order
            transcription = self.pending_transcriptions.pop(0)
            
            # Format timestamp
            timestamp = transcription['timestamp'].strftime("%H:%M:%S")
            speaker = transcription['speaker']
            text = transcription['text']
            
            # Add to meeting log
            entry = {
                'timestamp': timestamp,
                'speaker': speaker,
                'text': text,
                'user_id': transcription['user_id']
            }
            self.meeting["log"].append(entry)
            
            # Save meetings after adding new transcription data
            save_meetings()
            
            # Send to channel
            channel = self.meeting["channel"]
            if channel:
                try:
                    if self.stopped:
                        return
                    await channel.send(f"[{timestamp}] **{speaker}**: {text}")
                except discord.errors.NotFound:
                    logger.warning("Text channel not found, meeting may have ended")
                except Exception as e:
                    logger.error(f"Error sending message to channel: {e}")
            
            # Small delay to ensure proper ordering in Discord
            await asyncio.sleep(0.1)
    
    def transcribe_audio(self, wav_buffer):
        """Transcribe audio using Whisper model with language support."""
        try:
            
            language_code = None if self.language == "auto" else self.language
            
            segments, info = model.transcribe(
                wav_buffer, 
                language=language_code,
                vad_filter=False,
                vad_parameters=dict(min_silence_duration_ms=500),
                word_timestamps=True,
                condition_on_previous_text=False,
                compression_ratio_threshold=2.4,
                log_prob_threshold=-1.0,
                no_speech_threshold=0.6,
                temperature=0.2,
                beam_size=3
            )
            
            text_segments = []
            for segment in segments:
                
                if segment.text.strip() and len(segment.text.strip()) > 1:
                    text_segments.append(segment.text.strip())
            
            return " ".join(text_segments)
        except Exception as e:
            logger.error(f"Transcription error: {e}")
            return ""


async def get_meeting_context_from_channel(channel):
    """Get formatted meeting context from channel message history."""
    try:
        # Check if this is a meeting transcription channel
        if not (channel.name.startswith('meeting-transcription-') or 
                ('transcription' in channel.topic.lower() if channel.topic else False)):
            return None
        
        messages = []
        
        async for message in channel.history(limit=500, oldest_first=True):
            # Skip bot system messages and commands
            if message.author.bot and not message.content.startswith('['):
                continue
            if message.content.startswith('!'):
                continue
                
            
            if message.content.startswith('[') and '**' in message.content and '**:' in message.content:
                messages.append(message.content)
        
        if not messages:
            return None
        
        
        context = "Meeting Transcript:\n"
        context += f"Channel: {channel.name}\n"
        context += f"Total Messages: {len(messages)}\n"
        context += "="*50 + "\n\n"
        
        
        for msg in messages:
            context += f"{msg}\n"
        
        return context
        
    except Exception as e:
        logger.error(f"Error getting meeting context from channel: {e}")
        return None

async def get_meeting_context(guild_id):
    """Get formatted meeting context for OpenAI (for active meetings)."""
    if guild_id not in active_meetings:
        return None
    
    meeting = active_meetings[guild_id]
    if not meeting["log"]:
        return None
    
   
    context = "Meeting Transcript:\n"
    context += f"Started: {meeting['start_time'].strftime('%Y-%m-%d %H:%M:%S')}\n"
    context += f"Language: {meeting.get('language_display', 'Auto-detect')}\n"
    context += "="*50 + "\n\n"
    
    for entry in meeting["log"]:
        context += f"[{entry['timestamp']}] {entry['speaker']}: {entry['text']}\n"
    
    return context

async def ask_openai_about_meeting(transcript_context, question):
    """Ask OpenAI about the meeting content."""
    try:
        system_prompt = """You are a helpful assistant that answers questions about meeting transcripts. 
        You should only answer questions based on the provided transcript content. 
        If the information is not available in the transcript, politely say so.
        Be concise and accurate in your responses.
        When referencing what someone said, include their name and approximate time if available."""
        
        user_prompt = f"""Here is the meeting transcript:

{transcript_context}

Question: {question}

Please answer based only on the information provided in the transcript above."""

        response = await openai_client.chat.completions.create(
            model="gpt-3.5-turbo",  
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=500,
            temperature=0.3
        )
        
        return response.choices[0].message.content.strip()
    
    except Exception as e:
        logger.error(f"OpenAI API error: {e}")
        return f"Sorry, I encountered an error while processing your question: {str(e)}"


async def create_transcript_document(meeting_log, format_type="docx"):
    """Create a transcript document in the specified format."""
    if format_type == "docx":
        return await create_word_document(meeting_log)
    elif format_type == "pdf":
        return await create_pdf_document(meeting_log)
    else:
        return await create_text_document(meeting_log)

async def create_word_document(meeting_log):
    """Create a Word document transcript."""
    doc = Document()
    doc.add_heading("Meeting Transcript", 0)
    doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("=" * 50)
    
    for entry in meeting_log:
        doc.add_paragraph(f"[{entry['timestamp']}] {entry['speaker']}: {entry['text']}")
    
    filename = f"transcript_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename

async def create_pdf_document(meeting_log):
    """Create a PDF document transcript."""
    filename = f"transcript_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    
    story.append(Paragraph("Meeting Transcript", styles['Title']))
    story.append(Spacer(1, 20))
    story.append(Paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    story.append(Spacer(1, 20))
    
    
    for entry in meeting_log:
        line = f"[{entry['timestamp']}] {entry['speaker']}: {entry['text']}"
        story.append(Paragraph(line, styles['Normal']))
        story.append(Spacer(1, 6))
    
    doc.build(story)
    return filename

async def create_text_document(meeting_log):
    """Create a text file transcript."""
    filename = f"transcript_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    
    with open(filename, 'w', encoding='utf-8') as f:
        f.write("Meeting Transcript\n")
        f.write("=" * 50 + "\n")
        f.write(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        
        for entry in meeting_log:
            f.write(f"[{entry['timestamp']}] {entry['speaker']}: {entry['text']}\n")
    
    return filename


@bot.event
async def on_ready():
    """Called when the bot is ready."""
    logger.info(f'{bot.user} has connected to Discord!')
    print(f'{bot.user} has connected to Discord!')
    
    # Load meetings from persistent storage
    load_meetings()
    
    # Restore meeting channels for active meetings
    await restore_meeting_channels()
    
    print(f'Bot is ready! Logged in as {bot.user.name}')
    print(f'Bot ID: {bot.user.id}')
    print(f'Connected to {len(bot.guilds)} guilds')

async def restore_meeting_channels():
    """Restore meeting channels for active meetings after bot restart."""
    for guild_id, meeting in list(active_meetings.items()):
        try:
            guild = bot.get_guild(guild_id)
            if not guild:
                logger.warning(f"Guild {guild_id} not found, removing meeting")
                active_meetings.pop(guild_id)
                continue
            
            # Try to restore the channel
            if meeting.get("channel_id"):
                channel = guild.get_channel(meeting["channel_id"])
                if channel:
                    meeting["channel"] = channel
                    logger.info(f"Restored channel for guild {guild_id}")
                else:
                    # Channel was deleted, create a new one
                    channel_name = f"meeting-transcription-{meeting['start_time'].strftime('%m%d-%H%M')}"
                    overwrites = {
                        guild.default_role: discord.PermissionOverwrite(read_messages=True, send_messages=False),
                        guild.me: discord.PermissionOverwrite(read_messages=True, send_messages=True)
                    }
                    
                    try:
                        new_channel = await guild.create_text_channel(
                            channel_name,
                            overwrites=overwrites,
                            topic=f"Restored transcription channel | Language: {meeting['language_display']}"
                        )
                        meeting["channel"] = new_channel
                        meeting["channel_id"] = new_channel.id
                        logger.info(f"Created new channel for restored meeting in guild {guild_id}")
                    except Exception as e:
                        logger.error(f"Failed to create new channel for guild {guild_id}: {e}")
                        active_meetings.pop(guild_id)
                        continue
            
            # Send restoration message
            if meeting["channel"]:
                await meeting["channel"].send(
                    f"🔄 **Meeting Restored**\n"
                    f"This meeting was restored after a bot restart.\n"
                    f"Language: {meeting['language_display']}\n"
                    f"Started: {meeting['start_time'].strftime('%Y-%m-%d %H:%M:%S')}\n"
                    f"💡 Use `!end_meeting` to generate the final transcript!\n"
                    + "="*50
                )
                
        except Exception as e:
            logger.error(f"Error restoring meeting for guild {guild_id}: {e}")
            active_meetings.pop(guild_id)

@bot.event
async def on_voice_state_update(member, before, after):
    """Handle voice state changes to clean up meetings when everyone leaves."""
    
    for guild_id, meeting in list(active_meetings.items()):
        if meeting["vc"] and meeting["vc"].channel:
           
            members_in_channel = [m for m in meeting["vc"].channel.members if not m.bot]
            
            
            if len(members_in_channel) == 0:
                logger.info(f"Auto-ending meeting in guild {guild_id} - no members left")
                await auto_end_meeting(guild_id)

async def auto_end_meeting(guild_id):
    """Automatically end a meeting when all participants leave."""
    if guild_id not in active_meetings:
        return
    
    meeting = active_meetings.pop(guild_id)
    
    try:
        # Clean up the sink
        if hasattr(meeting.get("sink"), 'cleanup'):
            meeting["sink"].cleanup()
        
        # Stop voice connection
        if meeting["vc"]:
            await meeting["vc"].disconnect()
        
        # Generate and send transcript if there are entries
        if meeting["log"]:
            filename = await create_transcript_document(meeting["log"], "pdf")
            channel = meeting["channel"]
            
            if channel:
                try:
                    with open(filename, 'rb') as f:
                        file = discord.File(f, filename)
                        await channel.send("📄 Meeting ended automatically. Here's the transcript:", file=file)
                    
                    # Clean up file
                    os.remove(filename)
                except Exception as e:
                    logger.error(f"Error sending auto-generated transcript: {e}")
            else:
                logger.warning(f"Meeting channel not available for guild {guild_id}, cannot send auto-generated transcript")
        
    except Exception as e:
        logger.error(f"Error in auto_end_meeting: {e}")

# ==== Commands ====
@bot.command(name='start_meeting', aliases=['start', 'begin_meeting'])
async def start_meeting(ctx, *, language:str|None=None):
    """Start a new meeting with transcription in specified language.
    
    Usage: !start_meeting [language]
    Examples: 
        !start_meeting
        !start_meeting english
        !start_meeting arabic
        !start_meeting auto
    """
    # Check if user is in a voice channel
    if ctx.author.voice is None:
        await ctx.send("❌ You must be in a voice channel to start a meeting.")
        return
    
    guild_id = ctx.guild.id
    
    # Check if a meeting is already active
    if guild_id in active_meetings:
        await ctx.send("❌ A meeting is already active in this server. Use `!end_meeting` to stop the current meeting.")
        return
    
    # Process language input
    language_code = get_language_code(language)
    language_display = get_language_display_name(language_code)
    
    guild = ctx.guild
    
    try:
        # Create a dedicated text channel for transcription
        channel_name = f"meeting-transcription-{datetime.datetime.now().strftime('%m%d-%H%M')}"
        overwrites = {
            guild.default_role: discord.PermissionOverwrite(read_messages=True, send_messages=False),
            guild.me: discord.PermissionOverwrite(read_messages=True, send_messages=True)
        }
        
        text_channel = await guild.create_text_channel(
            channel_name, 
            overwrites=overwrites,
            topic=f"Live transcription for meeting started by {ctx.author.display_name} | Language: {language_display}"
        )
        
        await text_channel.send(
            f"🎙️ **Meeting Transcription Started**\n"
            f"Started by: {ctx.author.mention}\n"
            f"Voice Channel: {ctx.author.voice.channel.name}\n"
            f"Language: {language_display}\n"
            f"💡 Use `!ask <question>` to ask about the meeting content!\n"
            + "="*50
        )
        
    except discord.Forbidden:
        # Fallback if bot cannot create channels
        text_channel = ctx.channel
        await ctx.send(f"⚠️ Missing permissions to create channels. Using this channel for transcription.\n🌐 Language: {language_display}")
    
    # Connect to voice channel
    vc_channel = ctx.author.voice.channel
    try:
        voice_client = await vc_channel.connect(cls=voice_recv.VoiceRecvClient)
    except Exception as e:
        await ctx.send(f"❌ Failed to connect to voice channel: {e}")
        return
    
    # Create meeting record
    meeting = {
        "channel": text_channel,
        "log": [],
        "vc": voice_client,
        "start_time": datetime.datetime.now(),
        "started_by": ctx.author.id,
        "language": language_code,
        "language_display": language_display
    }
    
    active_meetings[guild_id] = meeting
    save_meetings() # Save meeting after it's created
    
    # Set up audio transcription with language support
    sink = TranscriptionSink(meeting, bot, language_code)
    meeting["sink"] = sink  # Store reference for cleanup
    def after_listen(err: Exception | None):
    # O callback after pode ocorrer fora do loop async.
    # Então agendamos uma corrotina no loop do bot.
        if err is None:
            return

        # Detecta o caso específico: OpusError corrupted stream
        if isinstance(err, discord.opus.OpusError) and "corrupted stream" in str(err).lower():
            loop = bot.loop
            if loop and not loop.is_closed():
                loop.call_soon_threadsafe(
                    lambda: asyncio.create_task(
                        fail_meeting_capture(
                            guild_id=guild_id,
                            meeting=meeting,
                            reason="OpusError: corrupted stream",
                            ctx=ctx
                        )
                    )
                )
            return

        # Outros erros: registra, e opcionalmente também encerra
        logger.error(f"Voice listen error (non-opus): {err}")

    # listen com fallback
    voice_client.listen(sink, after=after_listen)        
        
    await ctx.send(f"✅ Meeting started! Live transcriptions in **{language_display}** will appear in {text_channel.mention}")

@bot.command(name='end_meeting', aliases=['end', 'stop_meeting'])
async def end_meeting(ctx, format_type="pdf"):
    """End the current meeting and generate transcript."""
    guild_id = ctx.guild.id
    
    if guild_id not in active_meetings:
        await ctx.send("❌ No active meeting found in this server.")
        return
    
    meeting = active_meetings.pop(guild_id)
    save_meetings() # Save meeting before it's removed
    # define canal cedo para evitar UnboundLocalError
    channel = meeting.get("channel") or ctx.channel
    
    try:
        # Clean up the sink
        if hasattr(meeting.get("sink"), 'cleanup'):
            meeting["sink"].cleanup()
        
        # Disconnect from voice
        if meeting["vc"]:
            await meeting["vc"].disconnect()
        
        mix_path = meeting.get("mix_audio_path")
        if mix_path and os.path.exists(mix_path):
            await channel.send(
                f"🎧 **Áudio completo (mix) salvo localmente**:\n`{mix_path}`\n"
            )
        else:
            await channel.send("⚠️ Não encontrei o arquivo de áudio mix no disco.")
       
        # Generate transcript document
        if meeting["log"]:
            # Validate format type
            if format_type.lower() not in ["pdf", "docx", "txt"]:
                format_type = "pdf"
            
            filename = await create_transcript_document(meeting["log"], format_type.lower())
            

            if not channel:
                # If meeting channel is not available (e.g., after bot restart), use the command channel
                channel = ctx.channel
                logger.info(f"Meeting channel not available, using command channel {ctx.channel.name} for transcript")
            
            duration = datetime.datetime.now() - meeting["start_time"]
            
            embed = discord.Embed(
                title="📄 Meeting Ended",
                description=f"Meeting duration: {str(duration).split('.')[0]}",
                color=0x00ff00
            )
            embed.add_field(name="Total Messages", value=len(meeting["log"]), inline=True)
            embed.add_field(name="Format", value=format_type.upper(), inline=True)
            
            try:
                with open(filename, 'rb') as f:
                    file = discord.File(f, filename)
                    await channel.send(embed=embed, file=file)
                
                # Clean up file
                os.remove(filename)
                await ctx.send("✅ Meeting ended successfully! Transcript has been generated.")
                
            except discord.Forbidden:
                logger.error(f"Bot lacks permission to send files in channel {channel.name}")
                await ctx.send("❌ Bot lacks permission to send files in the transcription channel. Please check bot permissions.")
            except discord.HTTPException as e:
                logger.error(f"HTTP error sending transcript file: {e}")
                await ctx.send(f"❌ Error sending transcript file: {e}")
            except Exception as e:
                logger.error(f"Error sending transcript file: {e}")
                await ctx.send("✅ Meeting ended, but there was an error sending the transcript file.")
        else:
            await ctx.send("✅ Meeting ended. No transcription data was recorded.")
            
    except Exception as e:
        logger.error(f"Error ending meeting: {e}")
        await ctx.send("❌ Error ending meeting. Please try again.")

@bot.command(name='ask')
async def ask_meeting_question(ctx, *, question):
    """Ask a question about the meeting transcript in the current channel.
    
    Usage: !ask <your question>
    Example: !ask What did John say about the budget?
    
    This command works in:
    - Active meeting transcription channels
    - Completed meeting transcription channels (using message history)
    """
    if not question or len(question.strip()) == 0:
        await ctx.send("❌ Please provide a question to ask about the meeting.")
        return
    
    # First try to get context from active meeting
    guild_id = ctx.guild.id
    transcript_context = None
    
    if guild_id in active_meetings:
        active_meeting = active_meetings[guild_id]
        if active_meeting["channel"].id == ctx.channel.id:
            # User is asking in the active meeting channel
            transcript_context = await get_meeting_context(guild_id)
    
    # If no active meeting context, try to get from channel history
    if not transcript_context:
        transcript_context = await get_meeting_context_from_channel(ctx.channel)
    
    if not transcript_context:
        await ctx.send(
            "❌ No meeting transcript found in this channel.\n"
            "💡 Use this command in:\n"
            "• An active meeting transcription channel\n" 
            "• A completed meeting transcription channel\n"
            "• Start a meeting with `!start_meeting` first"
        )
        return
    
    # Check if transcript is too long and truncate if needed (to avoid token limits)
    max_length = 12000  # Approximately 3000 tokens
    if len(transcript_context) > max_length:
        lines = transcript_context.split('\n')
        header_lines = lines[:5]  # Keep header info
        transcript_lines = [line for line in lines[5:] if line.startswith('[')]
        
        # Take recent messages if truncation is needed
        if len('\n'.join(transcript_lines)) > max_length - 500:
            transcript_lines = transcript_lines[-100:]  # Last 100 messages
            header_lines.append("⚠️ [Transcript truncated to recent messages due to length]")
        
        transcript_context = '\n'.join(header_lines + transcript_lines)
    
    # Show typing indicator
    async with ctx.typing():
        try:
            # Get answer from OpenAI
            answer = await ask_openai_about_meeting(transcript_context, question)
            
            # Create embed for the response
            embed = discord.Embed(
                title="🤖 Meeting Q&A",
                color=0x00ff99
            )
            embed.add_field(name="❓ Question", value=question, inline=False)
            embed.add_field(name="💬 Answer", value=answer, inline=False)
            
            # Add context info
            if guild_id in active_meetings and active_meetings[guild_id]["channel"].id == ctx.channel.id:
                embed.set_footer(text="Based on active meeting transcript")
            else:
                embed.set_footer(text="Based on channel message history")
            
            await ctx.send(embed=embed)
            
        except Exception as e:
            logger.error(f"Error in ask command: {e}")
            await ctx.send("❌ Sorry, I encountered an error while processing your question. Please try again.")

@bot.command(name='meeting_status', aliases=['status', 'meeting_info'])
async def meeting_status(ctx):
    """Check current meeting status."""
    guild_id = ctx.guild.id
    
    if guild_id in active_meetings:
        active_meeting = active_meetings[guild_id]
        start_time = active_meeting["start_time"]
        duration = datetime.datetime.now() - start_time
        duration_str = str(duration).split('.')[0]  # Remove microseconds
        
        embed = discord.Embed(
            title="🎙️ Active Meeting Status",
            color=0x00ff00
        )
        
        embed.add_field(
            name="Status", 
            value="🟢 **ACTIVE**", 
            inline=True
        )
        embed.add_field(
            name="Duration", 
            value=f"⏱️ {duration_str}", 
            inline=True
        )
        embed.add_field(
            name="Language", 
            value=f"🌐 {active_meeting['language_display']}", 
            inline=True
        )
        embed.add_field(
            name="Started By", 
            value=f"👤 <@{active_meeting['started_by']}>", 
            inline=True
        )
        embed.add_field(
            name="Transcription Channel", 
            value=f"📝 {active_meeting['channel'].mention if active_meeting['channel'] else 'Not available'}", 
            inline=True
        )
        embed.add_field(
            name="Messages Transcribed", 
            value=f"💬 {len(active_meeting['log'])}", 
            inline=True
        )
        
        embed.set_footer(text="Use !end_meeting to finish and generate transcript")
        
        await ctx.send(embed=embed)
    else:
        embed = discord.Embed(
            title="🎙️ Meeting Status",
            description="No active meeting in this server.",
            color=0xff0000
        )
        embed.add_field(
            name="To start a meeting", 
            value="• Join a voice channel\n• Use `!start_meeting [language]`\n• Examples: `!start_meeting english`, `!start_meeting arabic`", 
            inline=False
        )
        await ctx.send(embed=embed)

@bot.command(name='restore_meeting', aliases=['restore', 'recover_meeting'])
async def restore_meeting(ctx):
    """Manually restore a meeting if it was lost due to bot restart."""
    guild_id = ctx.guild.id
    
    if guild_id in active_meetings:
        await ctx.send("✅ A meeting is already active in this server.")
        return
    
    # Check if there's a saved meeting file
    if os.path.exists("meetings.pkl"):
        try:
            with open("meetings.pkl", "rb") as f:
                saved_meetings = pickle.load(f)
            
            if guild_id in saved_meetings:
                saved_meeting = saved_meetings[guild_id]
                
                # Check if meeting is still valid (within last 24 hours)
                if (datetime.datetime.now() - saved_meeting["start_time"]).total_seconds() < 86400:
                    # Restore the meeting
                    active_meetings[guild_id] = {
                        "channel": None,
                        "log": saved_meeting["log"],
                        "vc": None,
                        "start_time": saved_meeting["start_time"],
                        "started_by": saved_meeting["started_by"],
                        "language": saved_meeting["language"],
                        "language_display": saved_meeting["language_display"],
                        "sink": None
                    }
                    
                    # Try to restore the channel
                    guild = ctx.guild
                    if saved_meeting.get("channel_id"):
                        channel = guild.get_channel(saved_meeting["channel_id"])
                        if channel:
                            active_meetings[guild_id]["channel"] = channel
                            active_meetings[guild_id]["channel_id"] = channel.id
                            await ctx.send(f"✅ Meeting restored! Channel: {channel.mention}")
                        else:
                            # Channel was deleted, create a new one
                            channel_name = f"meeting-transcription-{saved_meeting['start_time'].strftime('%m%d-%H%M')}"
                            overwrites = {
                                guild.default_role: discord.PermissionOverwrite(read_messages=True, send_messages=False),
                                guild.me: discord.PermissionOverwrite(read_messages=True, send_messages=True)
                            }
                            
                            try:
                                new_channel = await guild.create_text_channel(
                                    channel_name,
                                    overwrites=overwrites,
                                    topic=f"Restored transcription channel | Language: {saved_meeting['language_display']}"
                                )
                                active_meetings[guild_id]["channel"] = new_channel
                                active_meetings[guild_id]["channel_id"] = new_channel.id
                                await ctx.send(f"✅ Meeting restored! New channel created: {new_channel.mention}")
                            except Exception as e:
                                logger.error(f"Failed to create new channel for guild {guild_id}: {e}")
                                await ctx.send("⚠️ Meeting restored but failed to create new channel. Use `!end_meeting` in any channel to generate transcript.")
                    else:
                        await ctx.send("✅ Meeting restored! Use `!end_meeting` to generate transcript.")
                    
                    save_meetings()
                else:
                    await ctx.send("❌ Saved meeting has expired (older than 24 hours).")
            else:
                await ctx.send("❌ No saved meeting found for this server.")
        except Exception as e:
            await ctx.send(f"❌ Error restoring meeting: {e}")
    else:
        await ctx.send("❌ No saved meetings file found.")

@bot.command(name='fix_channel', aliases=['fix_ch', 'repair_channel'])
async def fix_channel(ctx):
    """Fix channel reference for the current meeting if it's broken."""
    guild_id = ctx.guild.id
    
    if guild_id not in active_meetings:
        await ctx.send("❌ No active meeting found in this server.")
        return
    
    meeting = active_meetings[guild_id]
    
    if meeting["channel"] and meeting["channel"].id == ctx.channel.id:
        await ctx.send("✅ Channel reference is already correct for this meeting.")
        return
    
    # Update the channel reference to the current channel
    meeting["channel"] = ctx.channel
    meeting["channel_id"] = ctx.channel.id
    
    await ctx.send(
        f"✅ Channel reference fixed! This channel ({ctx.channel.mention}) is now the transcription channel.\n"
        f"💡 You can now use `!end_meeting` to generate the transcript."
    )
    
    # Save the updated meeting
    save_meetings()

@bot.command(name='low_latency_mode', aliases=['lowlatency', 'fast_mode', 'realtime'])
async def low_latency_mode(ctx, mode: str = None):
    """Enable or disable low-latency transcription mode.
    
    Usage: 
        !low_latency_mode - Show current mode
        !low_latency_mode on - Enable low-latency mode (faster but more CPU)
        !low_latency_mode off - Disable low-latency mode (slower but less CPU)
        !low_latency_mode ultra - Enable ultra-low latency (fastest, highest CPU)
    """
    guild_id = ctx.guild.id
    
    if guild_id not in active_meetings:
        await ctx.send("❌ No active meeting found in this server.")
        return
    
    meeting = active_meetings[guild_id]
    sink = meeting.get("sink")
    
    if not sink:
        await ctx.send("❌ Transcription sink not available.")
        return
    
    if mode is None:
        # Show current mode
        embed = discord.Embed(
            title="⚡ Low Latency Mode Status",
            description="Current transcription latency configuration",
            color=0x00ff00 if sink.real_time_mode else 0xff0000
        )
        
        status = "🟢 **ENABLED**" if sink.real_time_mode else "🔴 **DISABLED**"
        embed.add_field(name="Status", value=status, inline=True)
        
        embed.add_field(
            name="Buffer Duration", 
            value=f"⏱️ {sink.buffer_duration}s", 
            inline=True
        )
        embed.add_field(
            name="Force Process Interval", 
            value=f"⏱️ {sink.force_process_interval}s", 
            inline=True
        )
        
        embed.add_field(
            name="Current Latency", 
            value=f"📊 ~{sink.buffer_duration + 0.5:.1f}s total delay", 
            inline=True
        )
        
        embed.add_field(
            name="💡 Commands", 
            value="• `!low_latency_mode on` - Enable fast mode\n• `!low_latency_mode ultra` - Enable ultra-fast mode\n• `!low_latency_mode off` - Disable fast mode", 
            inline=False
        )
        
        await ctx.send(embed=embed)
        return
    
    # Change mode
    if mode.lower() in ["on", "true", "1", "fast"]:
        sink.real_time_mode = True
        sink.buffer_duration = 2.0
        sink.min_buffer_size = 8000
        sink.max_buffer_size = 24000
        sink.force_process_interval = 1.5
        await ctx.send("✅ Low-latency mode enabled! Buffer duration: 2.0s (~2.5s total delay)")
        
    elif mode.lower() in ["ultra", "fastest", "minimal"]:
        sink.real_time_mode = True
        sink.buffer_duration = 1.0
        sink.min_buffer_size = 4000
        sink.max_buffer_size = 16000
        sink.force_process_interval = 0.8
        await ctx.send("🚀 Ultra-low latency mode enabled! Buffer duration: 1.0s (~1.5s total delay)")
        
    elif mode.lower() in ["off", "false", "0", "slow"]:
        sink.real_time_mode = False
        sink.buffer_duration = 5.0
        sink.min_buffer_size = 16000
        sink.max_buffer_size = 48000
        sink.force_process_interval = 3.0
        await ctx.send("🐌 Low-latency mode disabled. Buffer duration: 5.0s (~5.5s total delay)")
        
    else:
        await ctx.send("❌ Invalid mode. Use: `on`, `ultra`, or `off`")

@bot.command(name='streaming_mode', aliases=['stream', 'partial_transcriptions'])
async def streaming_mode(ctx, mode: str = None):
    """Enable or disable streaming transcriptions for ultra-low perceived latency.
    
    Usage: 
        !streaming_mode - Show current mode
        !streaming_mode on - Enable streaming (shows partial transcriptions)
        !streaming_mode off - Disable streaming (shows only complete transcriptions)
    """
    guild_id = ctx.guild.id
    
    if guild_id not in active_meetings:
        await ctx.send("❌ No active meeting found in this server.")
        return
    
    meeting = active_meetings[guild_id]
    sink = meeting.get("sink")
    
    if not sink:
        await ctx.send("❌ Transcription sink not available.")
        return
    
    if mode is None:
        # Show current mode
        embed = discord.Embed(
            title="🌊 Streaming Mode Status",
            description="Current transcription streaming configuration",
            color=0x00ff00 if sink.streaming_mode else 0xff0000
        )
        
        status = "🟢 **ENABLED**" if sink.streaming_mode else "🔴 **DISABLED**"
        embed.add_field(name="Status", value=status, inline=True)
        
        embed.add_field(
            name="Perceived Latency", 
            value=f"📊 {'~0.5s' if sink.streaming_mode else '~2.5s'} (with partial transcriptions)", 
            inline=True
        )
        
        embed.add_field(
            name="💡 How It Works", 
            value="Shows partial transcriptions as they're being processed, giving the illusion of instant response", 
            inline=False
        )
        
        embed.add_field(
            name="💡 Commands", 
            value="• `!streaming_mode on` - Enable streaming\n• `!streaming_mode off` - Disable streaming", 
            inline=False
        )
        
        await ctx.send(embed=embed)
        return
    
    # Change mode
    if mode.lower() in ["on", "true", "1", "enable"]:
        sink.streaming_mode = True
        await ctx.send("✅ Streaming mode enabled! You'll see partial transcriptions for ultra-low perceived latency.")
        
    elif mode.lower() in ["off", "false", "0", "disable"]:
        sink.streaming_mode = False
        await ctx.send("🔴 Streaming mode disabled. Only complete transcriptions will be shown.")
        
    else:
        await ctx.send("❌ Invalid mode. Use: `on` or `off`")

@bot.command(name='transcription_settings', aliases=['trans_settings', 'audio_settings'])
async def transcription_settings(ctx, setting: str = None, value: float = None):
    """View or adjust transcription settings for better accuracy and ordering.
    
    Usage: 
        !transcription_settings - Show current settings
        !transcription_settings buffer_duration 1.0 - Set buffer duration to 1 second
        !transcription_settings min_buffer_size 4000 - Set minimum buffer size
        !transcription_settings max_buffer_size 16000 - Set maximum buffer size
        !transcription_settings force_interval 0.8 - Set force processing interval
    """
    guild_id = ctx.guild.id
    
    if guild_id not in active_meetings:
        await ctx.send("❌ No active meeting found in this server.")
        return
    
    meeting = active_meetings[guild_id]
    sink = meeting.get("sink")
    
    if not sink:
        await ctx.send("❌ Transcription sink not available.")
        return
    
    if setting is None:
        # Show current settings
        embed = discord.Embed(
            title="🎙️ Transcription Settings",
            description="Current audio processing configuration",
            color=0x0099ff
        )
        
        embed.add_field(
            name="Buffer Duration", 
            value=f"⏱️ {sink.buffer_duration}s (time before processing audio)", 
            inline=True
        )
        embed.add_field(
            name="Min Buffer Size", 
            value=f"📊 {sink.min_buffer_size} bytes (minimum audio before processing)", 
            inline=True
        )
        embed.add_field(
            name="Max Buffer Size", 
            value=f"📊 {sink.max_buffer_size} bytes (force processing if exceeded)", 
            inline=True
        )
        
        embed.add_field(
            name="Force Process Interval", 
            value=f"⏱️ {sink.force_process_interval}s (background processing)", 
            inline=True
        )
        
        embed.add_field(
            name="Real-time Mode", 
            value=f"⚡ {'🟢 ON' if sink.real_time_mode else '🔴 OFF'}", 
            inline=True
        )
        
        embed.add_field(
            name="💡 Tips", 
            value="• Lower buffer duration = faster response but more CPU\n• Higher buffer size = better accuracy but longer delays\n• Use `!low_latency_mode ultra` for fastest response\n• Use `!transcription_settings <setting> <value>` to adjust", 
            inline=False
        )
        
        await ctx.send(embed=embed)
        return
    
    # Adjust settings
    if value is None:
        await ctx.send(f"❌ Please provide a value for {setting}. Example: `!transcription_settings {setting} 1.0`")
        return
    
    if setting == "buffer_duration":
        if 0.5 <= value <= 10.0:
            sink.buffer_duration = value
            await ctx.send(f"✅ Buffer duration set to {value}s")
        else:
            await ctx.send("❌ Buffer duration must be between 0.5 and 10.0 seconds")
    elif setting == "min_buffer_size":
        if 2000 <= value <= 32000:
            sink.min_buffer_size = int(value)
            await ctx.send(f"✅ Minimum buffer size set to {int(value)} bytes")
        else:
            await ctx.send("❌ Minimum buffer size must be between 2000 and 32000 bytes")
    elif setting == "max_buffer_size":
        if 8000 <= value <= 64000:
            sink.max_buffer_size = int(value)
            await ctx.send(f"✅ Maximum buffer size set to {int(value)} bytes")
        else:
            await ctx.send("❌ Maximum buffer size must be between 8000 and 64000 bytes")
    elif setting == "force_interval":
        if 0.5 <= value <= 5.0:
            sink.force_process_interval = value
            await ctx.send(f"✅ Force processing interval set to {value}s")
        else:
            await ctx.send("❌ Force processing interval must be between 0.5 and 5.0 seconds")
    else:
        await ctx.send(f"❌ Unknown setting '{setting}'. Available settings: buffer_duration, min_buffer_size, max_buffer_size, force_interval")

# ==== Error Handling ====
@bot.event
async def on_command_error(ctx, error):
    if isinstance(error, commands.CommandNotFound):
        return  # Ignore unknown commands
    elif isinstance(error, commands.MissingPermissions):
        await ctx.send("❌ You don't have permission to use this command.")
    elif isinstance(error, commands.MissingRequiredArgument):
        if ctx.command.name == "ask":
            await ctx.send("❌ Please provide a question. Usage: `!ask <your question>`")
        else:
            await ctx.send(f"❌ Missing required argument for command `{ctx.command.name}`")
    else:
        logger.error(f"Command error: {error}")
        await ctx.send(f"❌ An error occurred: {str(error)}")

# ==== Help Command ====
@bot.command(name='meeting_help', aliases=['help_meeting'])
async def meeting_help(ctx):
    """Show help for meeting commands."""
    embed = discord.Embed(
        title="🎙️ Meeting Transcription Bot Help",
        description="Commands for managing meeting transcriptions and Q&A",
        color=0x0099ff
    )
    
    embed.add_field(
        name="!start_meeting [language]", 
        value="Start meeting with live transcription in specified language\nExamples: `!start_meeting arabic`, `!start_meeting english`", 
        inline=False
    )
    embed.add_field(
        name="!end_meeting [format]", 
        value="End meeting and generate transcript (format: pdf, docx, txt)", 
        inline=False
    )
    embed.add_field(
        name="!ask <question>", 
        value="Ask a question about meeting transcript in current channel\nWorks during and after meetings in transcription channels\nExample: `!ask What did Sarah say about the deadline?`", 
        inline=False
    )
    embed.add_field(
        name="!meeting_status", 
        value="Check current meeting status", 
        inline=False
    )
    embed.add_field(
        name="!restore_meeting", 
        value="Restore a meeting if it was lost due to bot restart", 
        inline=False
    )
    embed.add_field(
        name="!fix_channel", 
        value="Fix channel reference for current meeting if broken", 
        inline=False
    )
    embed.add_field(
        name="!transcription_settings", 
        value="View or adjust audio processing settings for better accuracy", 
        inline=False
    )
    embed.add_field(
        name="!low_latency_mode", 
        value="Enable or disable low-latency transcription mode", 
        inline=False
    )
    embed.add_field(
        name="!streaming_mode", 
        value="Enable or disable streaming transcriptions for ultra-low perceived latency", 
        inline=False
    )
    embed.add_field(
        name="!languages", 
        value="Show supported languages for transcription", 
        inline=False
    )
    embed.add_field(
        name="!meeting_help", 
        value="Show this help message", 
        inline=False
    )
    
    embed.set_footer(text="Bot requires permissions: Connect, Speak, Use Voice Activity")
    
    await ctx.send(embed=embed)

@bot.command(name='languages', aliases=['supported_languages', 'lang'])
async def show_languages(ctx):
    """Show supported languages for transcription."""
    embed = discord.Embed(
        title="🌐 Supported Languages",
        description="Languages available for meeting transcription",
        color=0x00ff99
    )
    
    # Group languages by region for better organization
    languages_by_region = {
        "🌍 European": ["english", "spanish", "french", "german", "italian", "portuguese", "russian", "dutch", "swedish", "norwegian", "danish", "finnish", "polish", "czech", "hungarian", "greek"],
        "🏛️ Middle Eastern": ["arabic", "hebrew", "persian", "turkish", "urdu"],
        "🏮 Asian": ["chinese", "japanese", "korean", "hindi", "thai", "vietnamese", "indonesian", "malay", "tamil", "bengali"],
        "🤖 Special": ["auto (automatic detection)"]
    }
    
    for region, langs in languages_by_region.items():
        lang_list = ", ".join(f"`{lang}`" for lang in langs)
        embed.add_field(name=region, value=lang_list, inline=False)
    
    embed.add_field(
        name="📝 Usage Examples",
        value="`!start_meeting arabic`\n`!start_meeting english`\n`!start_meeting auto`",
        inline=False
    )
    
    embed.set_footer(text="Language detection is automatic if not specified")
    
    await ctx.send(embed=embed)

# ==== Run Bot ====
if __name__ == "__main__":
    # Check for required environment variables
    discord_token = os.getenv('DISCORD_TOKEN')
    openai_api_key = os.getenv('OPENAI_API_KEY')
    
    if not discord_token:
        print("ERROR: DISCORD_TOKEN not found in environment variables!")
        print("Please create a .env file with your bot token:")
        print("DISCORD_TOKEN=your_token_here")
        exit(1)
    
    if not openai_api_key:
        print("WARNING: OPENAI_API_KEY not found in environment variables!")
        print("The !ask command will not work without OpenAI API key.")
        print("Add to your .env file: OPENAI_API_KEY=your_api_key_here")
    
    print("🤖 Starting Discord Meeting Bot...")
    print("📋 Features enabled:")
    print("   ✅ Live transcription")
    print("   ✅ Multi-language support") 
    print("   ✅ Document generation (PDF, DOCX, TXT)")
    if openai_api_key:
        print("   ✅ AI Q&A about meetings")
    else:
        print("   ❌ AI Q&A (missing OpenAI API key)")
    
    print("🚀 Bot starting...")
    try:
        bot.run(discord_token)
    except Exception as e:
        print(f"❌ Error starting bot: {e}")
        exit(1)